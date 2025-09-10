# app_including_Activities (old).py
# Grant Application Portal – Logframe + Workplan + Budget
import streamlit as st
import pandas as pd
from io import BytesIO
from openpyxl import Workbook
import uuid
import base64
import os
import html, re
from html import escape
from datetime import datetime, date
import hashlib

# ---------------- Page config ----------------
st.set_page_config(page_title="Falcon Awards Application Portal", layout="wide")

# ---------------- Logo (top-right PNG, simple) ----------------
def add_logo(path="glide_logo.png", width_px=140):
    if not os.path.exists(path):
        return
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    st.markdown(
        f"""
        <div style="
            position: fixed;
            top: 40px; right: 900px;
            width: {width_px}px; z-index: 1000;">
            <img src="data:image/png;base64,{b64}" style="width:100%;height:auto;" alt="GLIDE Logo">
        </div>
        """,
        unsafe_allow_html=True,
    )

add_logo()

# ---------------- Helpers & State ----------------
def generate_id():
    return str(uuid.uuid4())[:8]

# app state
for key in ["impacts", "outcomes", "outputs", "kpis", "workplan", "budget"]:
    if key not in st.session_state:
        st.session_state[key] = []

# edit-state for inline editing
for key in ["edit_goal", "edit_outcome", "edit_output", "edit_kpi", "edit_activity", "edit_budget_row"]:
    if key not in st.session_state:
        st.session_state[key] = None

def _find_by_id(lst, _id):
    for i, x in enumerate(lst):
        if x.get("id") == _id:
            return i
    return None

def delete_cascade(*, goal_id=None, outcome_id=None, output_id=None):
    """Delete an item and its children in the Goal->Outcome->Output->KPI hierarchy."""
    if goal_id:
        for oc in [o for o in st.session_state.outcomes if o.get("parent_id") == goal_id]:
            delete_cascade(outcome_id=oc["id"])
        st.session_state.impacts = [g for g in st.session_state.impacts if g["id"] != goal_id]

    if outcome_id:
        for out in [o for o in st.session_state.outputs if o.get("parent_id") == outcome_id]:
            delete_cascade(output_id=out["id"])
        st.session_state.outcomes = [o for o in st.session_state.outcomes if o["id"] != outcome_id]

    if output_id:
        # delete KPIs under output
        st.session_state.kpis = [k for k in st.session_state.kpis if k.get("parent_id") != output_id]
        st.session_state.outputs = [o for o in st.session_state.outputs if o["id"] != output_id]

def render_editable_item(
    *,
    item: dict,
    list_name: str,
    edit_flag_key: str,
    view_md_func,
    fields=None,
    default_label="Name",
    on_delete=None,
    key_prefix="logframe"   # <— NEW
):
    c1, c2, c3 = st.columns([0.85, 0.07, 0.08])

    # unique widget id base
    wid = f"{key_prefix}_{list_name}_{item['id']}"

    if st.session_state.get(edit_flag_key) == item["id"]:
        new_values = {}
        if fields:
            for fkey, widget_func, label in fields:
                new_values[fkey] = widget_func(label, value=item.get(fkey, ""), key=f"{wid}_{fkey}")
        else:
            new_values["name"] = c1.text_input(default_label, value=item.get("name", ""), key=f"{wid}_name")

        if c2.button("💾", key=f"{wid}_save"):
            idx = _find_by_id(st.session_state[list_name], item["id"])
            if idx is not None:
                # Defensive: prevent users from saving names that include labels
                if "name" in new_values:
                    if list_name == "activities":
                        new_values["name"] = strip_label_prefix(new_values["name"], "Activity")
                    elif list_name == "kpis":
                        new_values["name"] = strip_label_prefix(new_values["name"], "KPI")

                for k, v in new_values.items():
                    st.session_state[list_name][idx][k] = v.strip() if isinstance(v, str) else v
            st.session_state[edit_flag_key] = None
            st.rerun()

        if c3.button("✖️", key=f"{wid}_cancel"):
            st.session_state[edit_flag_key] = None
            st.rerun()
    else:
        c1.markdown(view_md_func(item), unsafe_allow_html=True)
        if c2.button("✏️", key=f"{wid}_edit"):
            st.session_state[edit_flag_key] = item["id"]
            st.rerun()
        if c3.button("🗑️", key=f"{wid}_del"):
            if on_delete:
                on_delete()

from datetime import date

from datetime import date

def compute_numbers(include_activities: bool = False):
    """
    Preserve user/excel order.
    Outputs: numbered per Outcome in list order.
    KPIs:    numbered per Output in list order (as stored in st.session_state.kpis).
    Activities (if requested): numbered per Output in list order.
    """
    out_num, kpi_num = {}, {}
    outcomes = st.session_state.get("outcomes", [])
    outputs  = st.session_state.get("outputs", [])
    kpis     = st.session_state.get("kpis", [])

    # Outputs numbered per Outcome (list order)
    for oc in outcomes:
        oc_outs = [o for o in outputs if o.get("parent_id") == oc["id"]]  # list order preserved
        for i, out in enumerate(oc_outs, start=1):
            out_num[out["id"]] = f"{i}"

    # KPIs numbered per Output (list order)
    for out_id, n in out_num.items():
        p = 1
        for k in kpis:                       # iterate as-is → preserves entry/import order
            if k.get("parent_id") == out_id:
                kpi_num[k["id"]] = f"{n}.{p}"
                p += 1

    if not include_activities:
        return out_num, kpi_num

    # Activities numbered per Output (list order)
    act_num = {}
    workplan = st.session_state.get("workplan", [])
    for out_id, n in out_num.items():
        q = 1
        for a in workplan:
            if a.get("output_id") == out_id:
                act_num[a["id"]] = f"{n}.{q}"
                q += 1
    return out_num, kpi_num, act_num

def strip_label_prefix(text: str, kind: str) -> str:
    """
    Remove labels like 'Activity 1.2 — ' or 'KPI 1.2.3: ' from a string.
    Accepts separators '—', ':', or '-'.
    """
    if not isinstance(text, str):
        return text
    pat = rf'^\s*{kind}\s+\d+(?:\.\d+)*\s*[—:\-]\s*'
    return re.sub(pat, '', text).strip()

def parse_date_like(v):
    """Return a datetime.date or None from common date formats or existing date/datetime/pandas types."""
    if v is None:
        return None

    # Handle pandas NaT / NaN early
    try:
        import pandas as pd
        if pd.isna(v):
            return None
        if isinstance(v, pd.Timestamp):
            return v.date()
    except Exception:
        pass

    if isinstance(v, datetime):
        return v.date()
    if isinstance(v, date):
        return v

    # Strings
    s = str(v).strip()
    if not s or s.lower() in ("none", "nan", "nat"):
        return None

    # Try a bunch of common formats (with and without HH:MM:SS)
    fmts = (
        "%Y-%m-%d",
        "%Y-%m-%d %H:%M:%S",
        "%d/%m/%Y",
        "%d/%m/%Y %H:%M:%S",
        "%m/%d/%Y",
        "%m/%d/%Y %H:%M:%S",
        "%Y/%m/%d",
        "%Y/%m/%d %H:%M:%S",
        "%d/%b/%Y",            # 03/Sep/2025
        "%d/%b/%Y %H:%M:%S",
        "%d-%b-%Y",            # 03-Sep-2025
        "%d-%b-%Y %H:%M:%S",
    )
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue

    return None

def fmt_dd_mmm_yyyy(v):
    """Return 'DD/MMM/YYYY' (e.g., 03/Sep/2025) or '' if not set/parsable."""
    try:
        import pandas as pd
        if pd.isna(v):                   # handles NaT / NaN
            return ""
        if isinstance(v, pd.Timestamp):  # valid timestamp -> format directly
            return v.strftime("%d/%b/%Y")
    except Exception:
        pass

    d = parse_date_like(v)
    return d.strftime("%d/%b/%Y") if d else ""

def fmt_money(val) -> str:
    """Return number with thousands dot and 2 decimals, e.g., 1.234.567,89."""
    try:
        x = float(val)
    except (TypeError, ValueError):
        return ""
    # First format in US style, then swap separators
    s = f"{x:,.2f}"                  # -> 1,234,567.89
    return s.replace(",", "␟").replace(".", ",").replace("␟", ".")

def view_logframe_element(inner_html: str, kind: str = "output") -> str:
    """Wrap inner HTML in a styled card. kind: 'output' | 'kpi' (or others later)."""
    return f"<div class='lf-card lf-card--{kind}'>{inner_html}</div>"

def build_logframe_docx():
    # Lazy import so app loads even if package is missing
    try:
        from docx import Document
        from docx.shared import Cm, RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception:
        st.error("`python-docx` is required. In your venv run:\n  pip uninstall -y docx\n  pip install -U python-docx")
        raise

    PRIMARY_SHADE = "0A2F41"

    def _shade(cell, hex_fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tcPr.append(shd)

    def _repeat_header(row):
        trPr = row._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)

    def _set_cell_text(cell, text, *, bold=False, white=False, align_left=True):
        cell.text = ""
        p = cell.paragraphs[0]
        if align_left:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text or "")
        run.bold = bool(bold)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        try:
            rpr = run._element.rPr
            rpr.rFonts.set(qn("w:ascii"), "Calibri")
            rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
            rpr.rFonts.set(qn("w:cs"), "Calibri")
        except Exception:
            pass
        if white:
            run.font.color.rgb = RGBColor(255, 255, 255)

    def _add_run(p, text, bold=False):
        run = p.add_run(text or "")
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        return run

    # ---- Document setup: Portrait + standard margins
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, side, Cm(2.54))

    # ---- Data & numbering
    goal_text = (st.session_state.impacts[0]["name"] if st.session_state.get("impacts") else "")
    outcome_text = (st.session_state.outcomes[0]["name"] if st.session_state.get("outcomes") else "")
    out_nums, kpi_nums = compute_numbers()

    def _sort_by_num(label):
        if not label:
            return (9999,)
        try:
            return tuple(int(x) for x in str(label).split("."))
        except Exception:
            return (9999,)

    # ---- Helper: add a 2-row banner (label row shaded, content row unshaded)
    def _add_banner_block(label_text, content_text):
        t = doc.add_table(rows=2, cols=4)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Row 0: label
        r0 = t.rows[0]
        c0 = r0.cells[0]
        c0.merge(r0.cells[3])
        _shade(c0, PRIMARY_SHADE)
        _set_cell_text(c0, label_text.upper(), bold=True, white=True)
        # Row 1: content
        r1 = t.rows[1]
        c1 = r1.cells[0]
        c1.merge(r1.cells[3])
        _set_cell_text(c1, content_text or "")
        doc.add_paragraph("")

    # ---- GOAL & OUTCOME banners (keep these)
    if goal_text:
        _add_banner_block("GOAL", goal_text)
    if outcome_text:
        _add_banner_block("OUTCOME", outcome_text)

    # ---- ONE main table for all Outputs & KPIs (no per-output banners)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0]
    labels = ("Output", "KPI", "Means of Verification", "Key Assumptions")
    for i, lab in enumerate(labels):
        _set_cell_text(hdr.cells[i], lab, bold=True, white=True)
        _shade(hdr.cells[i], PRIMARY_SHADE)
    _repeat_header(hdr)

    outputs = st.session_state.get("outputs", [])
    outputs = sorted(outputs, key=lambda o: _sort_by_num(out_nums.get(o["id"], "")))

    for out in outputs:
        out_num = out_nums.get(out["id"], "")
        kpis = [k for k in st.session_state.get("kpis", []) if k.get("parent_id") == out["id"]]

        if not kpis:
            r = tbl.add_row()
            _set_cell_text(r.cells[0], f"Output {out_num} — {out.get('name','')}")
            _set_cell_text(r.cells[1], "—")
            _set_cell_text(r.cells[2], "—")
            _set_cell_text(r.cells[3], out.get("assumptions","") or "—")
            continue

        first = len(tbl.rows)  # first KPI row index for this output in the main table
        for k in kpis:
            r = tbl.add_row()
            k_lab = kpi_nums.get(k["id"], "")
            bp = (k.get("baseline", "") or "").strip()
            tg = (k.get("target", "") or "").strip()
            sd = fmt_dd_mmm_yyyy(k.get("start_date"))
            ed = fmt_dd_mmm_yyyy(k.get("end_date"))

            # Reset KPI cell
            kcell = r.cells[1]
            kcell.text = ""
            p = kcell.paragraphs[0]

            # Title line
            # Title
            title = f"KPI ({k_lab}) — {k.get('name', '')}" if k_lab else k.get('name', '')
            _add_run(p, title);
            p.add_run("\n")

            # Baseline
            if bp:
                _add_run(p, "Baseline: ", bold=True);
                _add_run(p, bp);
                p.add_run("\n")

            # Target
            if tg:
                _add_run(p, "Target: ", bold=True);
                _add_run(p, tg);
                p.add_run("\n")

            # Start / End (separate lines)
            if sd or ed:
                _add_run(p, "Start: ", bold=True);
                _add_run(p, sd or "—");
                p.add_run("\n")
                _add_run(p, "End: ", bold=True);
                _add_run(p, ed or "—")

            # MoV column (third col)
            mov_text = (k.get("mov") or "").strip() or "—"
            _set_cell_text(r.cells[2], mov_text)

        last = len(tbl.rows) - 1
        # Merge Output & Assumptions cells across the KPI block
        _set_cell_text(tbl.cell(first, 0), f"Output {out_num} — {out.get('name','')}")
        _set_cell_text(tbl.cell(first, 3), out.get("assumptions","") or "—")
        if last > first:
            tbl.cell(first, 0).merge(tbl.cell(last, 0))
            tbl.cell(first, 3).merge(tbl.cell(last, 3))

    # ---- Save to buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def view_activity_readonly(a, label, id_to_output, id_to_kpi):
    out_name = id_to_output.get(a.get("output_id"), "(unassigned)")
    kpis_txt = ", ".join(id_to_kpi.get(kid, "") for kid in (a.get("kpi_ids") or [])) or "—"
    sd = fmt_dd_mmm_yyyy(a.get("start")) or "—"
    ed = fmt_dd_mmm_yyyy(a.get("end"))   or "—"

    notes_html = ""
    if a.get("notes"):
        notes_html = f"<div class='lf-line'><b>Notes:</b> {escape(a.get('notes',''))}</div>"

    body = (
        f"<div class='lf-activity-title'>Activity {escape(label)} — {escape(a.get('name',''))}</div>"
        f"<div class='lf-line'><b>Output:</b> {escape(out_name)}</div>"
        f"<div class='lf-line'><b>Owner:</b> {escape(a.get('owner','') or '—')}</div>"
        f"<div class='lf-line'><b>Start date:</b> {sd} &nbsp;&nbsp;•&nbsp;&nbsp; <b>End date:</b> {ed}</div>"
        f"<div class='lf-line'><b>Status:</b> {escape(a.get('status','planned'))} &nbsp;&nbsp; "
        f"<b>% complete:</b> {int(a.get('progress',0))}%</div>"
        f"<div class='lf-line'><b>Linked KPIs:</b> {escape(kpis_txt)}</div>"
        f"{notes_html}"
    )

    return view_logframe_element(body, kind="activity")

def view_budget_item_card(row, id_to_output) -> str:
    """
    row: [OutputID, Item, Category, Unit, Qty, Unit Cost, Currency, Total]
    id_to_output: {output_id -> output name}
    """
    out_id, item, cat, unit, qty, uc, cur, tot = row
    out_name = id_to_output.get(out_id, "(unassigned)")
    # rows (label/value)
    lines = [
        ("Output", out_name),
        ("Item", item or "—"),
        ("Category", cat or "—"),
        ("Unit", unit or "—"),
        ("Qty", str(qty or 0)),
        ("Unit Cost", fmt_money(uc or 0)),
        ("Currency", cur or "—"),
        ("Total", fmt_money(tot or 0)),
    ]
    body = "".join(
        f"<div class='lf-line'><b>{escape(lbl)}:</b> {escape(val) if lbl not in ('Unit Cost','Total') else val}</div>"
        for (lbl, val) in lines
    )
    # reuse the blue style
    return f"<div class='lf-card lf-card--budget'>{body}</div>"

# ---------------- CSS for cards ----------------
def inject_logframe_css():
    st.markdown("""
    <style>
/* Base card */
.lf-card{
  margin: 14px 0 16px;
  padding: 14px 16px;
  border-radius: 12px;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
  position: relative;               /* required for the left accent bar */
}

/* ========= OUTPUT (green) ========= */
.lf-card--output{
  background: #E9F3E1;              /* pastel green fill */
  border: 1px solid #91A76A;        /* brand green border (#91A76A) */
}
.lf-card--output::before{
  content:"";
  position:absolute; top:0; left:0; bottom:0; width:8px;
  background:#6E8C4B;               /* darker green accent bar */
  border-top-left-radius:12px;
  border-bottom-left-radius:12px;
}

/* ========= KPI (orange) ========= */
.lf-card--kpi{
  background: #FFEBD6;          /* soft orange */
  border: 1px solid #F2B277;    /* warm orange border */
  border-radius: 12px;
  padding: 12px 16px;
  margin: 10px 0 12px;
  position: relative;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
}
.lf-card--kpi::before{
  content:"";
  position:absolute; top:0; left:0; bottom:0; width:6px;
  background:#DD7A1A;          /* stronger orange accent */
  border-top-left-radius:12px;
  border-bottom-left-radius:12px;
}

/* ========= Activity (yellow) ========= */
.lf-card--activity{
  background: #FFFBEA;          /* soft yellow */
  border: 1px solid #F6D58E;    /* warm yellow border */
  border-radius: 12px;
  padding: 12px 16px;
  margin: 10px 0 12px;
  position: relative;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
}
.lf-card--activity::before{
  content:"";
  position:absolute; top:0; left:0; bottom:0; width:6px;
  background:#F59E0B;          /* amber accent */
  border-top-left-radius:12px;
  border-bottom-left-radius:12px;
}

/* Headings & text bits */
.lf-out-header{ margin: 0; font-weight: 700; }
.lf-kpi-title{ font-weight: 600; margin-bottom: 6px; }

/* Assumptions list inside the green Output card */
.lf-card--output .lf-ass-heading{
  font-weight: 600;
  margin: 6px 0 4px;                 /* tighter heading spacing */
}

.lf-card--output .lf-ass-list{
  margin: 4px 0 0 1.15rem;           /* small left indent, no extra top gap */
  padding: 0;
  list-style: disc outside;
}

.lf-card--output .lf-ass-list li{
  margin: 0;                          /* remove extra gaps between items */
  padding: 0;
  font-style: italic;                 /* italics */
  font-size: 0.92rem;                 /* smaller text */
  line-height: 1.25;                  /* single/compact spacing */
}

/* Chips - Check if linked to payment*/
.chip {
  display:inline-block; padding: 2px 8px; border-radius: 999px;
  background:#eef2ff; font-size: .85rem; border:1px solid #e2e8ff;
}
.chip.green { background:#e6f9ee; border-color:#c6f0d8; color:#046c4e; }

/* Dot bullet in Output header */
.dot { font-size: 1.05rem; margin-right: .4rem; }

/* (optional) generic line styling used in KPI details */
.lf-line { margin: 2px 0; font-size: 0.95rem; color: #444; }

/* ========= BUDGET (blue) ========= */
.lf-card--budget{
  background: #E8F0FE;           /* light blue */
  border: 1px solid #90B4FE;
  border-radius: 12px;
  padding: 12px 16px;
  margin: 10px 0 12px;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
}
.lf-budget-total{ font-weight:700; margin-top:10px; }

.lf-card--budget .budget-table td:nth-child(1){
  max-width: 520px;              /* tune as needed */
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}   

/* --- Compact budget row (blue) --- */
.lf-budget-row{
  background:#E8F0FE;
  border:1px solid #90B4FE;
  border-radius:12px;
  padding:10px 12px;
  margin:8px 0;
  display:flex;
  align-items:center;
  gap:10px;
  box-shadow:0 1px 2px rgba(0,0,0,.03);
}
.lf-budget-cells{
  display:flex;
  gap:14px;
  flex-wrap:wrap;
  align-items:baseline;
  width:100%;
  font-size:.95rem;
  color:#1f2937;
}
.lf-budget-cell{
  white-space:nowrap;
}
.lf-budget-money{
  font-variant-numeric: tabular-nums;
  text-align:right;
}
.lf-budget-chip{
  background:#eef2ff;
  border:1px solid #dbe2ff;
  border-radius:999px;
  padding:2px 8px;
  display:inline-block;
}
.lf-subtotal{ font-weight:700; margin:6px 0 12px; text-align:right; }
.lf-grandtotal{ font-weight:800; margin-top:12px; text-align:right; }
    </style>
    """, unsafe_allow_html=True)

# ---------------- Tabs ----------------
tabs = st.tabs([
    "📘 Instructions",
    "🪪 Identification",
    "🧱 Logframe",
    "🗂️ Workplan",
    "💵 Budget",
    "📤 Export"
])

# ===== TAB 1: Instructions =====
tabs[0].markdown(
    """
# 📝 Welcome to the Falcon Awards Application Portal

Please complete each section of your application:

1. **Identification** - Basic project and investigator information, including title, PI details, institution, dates, and contact info
2. **Logframe** – Your project goals, results, and indicators  
3. **Workplan** – Planned activities and timelines  
4. **Budget** – Detailed costing for each activity

Once done, export your application as an Excel file.

"""
)
# **Contact us**: Anderson E. Stanciole | astanciole@glideae.org

# --- Resume from Excel ---
uploaded_file = tabs[0].file_uploader("Resume Previous Submission (Excel)", type="xlsx")
if uploaded_file is not None:
    try:
        # Build a stable signature for the uploaded file
        file_bytes = uploaded_file.getvalue()
        file_sig = hashlib.md5(file_bytes).hexdigest()

        # Only (re)load if this is a new file or changed content
        if st.session_state.get("_resume_file_sig") != file_sig:
            xls = pd.ExcelFile(BytesIO(file_bytes))

            # ---- RESET state containers
            st.session_state.impacts = []
            st.session_state.outcomes = []
            st.session_state.outputs = []
            st.session_state.kpis = []

            # (optional) clear edit flags so they won't point to old IDs
            for _f in ("edit_goal", "edit_outcome", "edit_output", "edit_kpi"):
                st.session_state[_f] = None

            # ---- Read Summary and build items (without trusting parent ids yet)
            summary_df = pd.read_excel(xls, sheet_name="Summary")

            goals_by_text = {}
            outcomes_by_text = {}
            pending_outcome_parent_ref = {}   # temp: outcome_id -> raw parent cell
            pending_output_parent_ref  = {}   # temp: output_id  -> raw parent cell

            def _clean_str(v):
                if pd.isna(v):
                    return ""
                return str(v).strip()

            for _, row in summary_df.iterrows():
                lvl    = _clean_str(row.get("Level", ""))
                text   = _clean_str(row.get("Text / Title", ""))
                parent = _clean_str(row.get("Parent ID", ""))  # may hold old UUID or a name (older files)

                if lvl.lower() == "goal":
                    gid = generate_id()
                    st.session_state.impacts.append({"id": gid, "level": "Goal", "name": text})
                    if text:
                        goals_by_text[text] = gid

                elif lvl.lower() == "outcome":
                    oid = generate_id()
                    st.session_state.outcomes.append({"id": oid, "level": "Outcome", "name": text, "parent_id": None})
                    pending_outcome_parent_ref[oid] = parent
                    if text:
                        outcomes_by_text[text] = oid

                elif lvl.lower() == "output":
                    pid = generate_id()
                    name_clean = strip_label_prefix(text, "Output") or "Output"
                    assumptions = _clean_str(row.get("Assumptions", ""))  # <-- read column if present
                    st.session_state.outputs.append({
                        "id": pid,
                        "level": "Output",
                        "name": name_clean,
                        "parent_id": None,
                        "assumptions": assumptions,  # <-- store it
                    })
                    pending_output_parent_ref[pid] = parent

            # ---- Resolve parents for Outcomes
            single_goal_id = st.session_state.impacts[0]["id"] if len(st.session_state.impacts) == 1 else None
            for oc in st.session_state.outcomes:
                raw = pending_outcome_parent_ref.get(oc["id"], "")
                if raw in goals_by_text:
                    oc["parent_id"] = goals_by_text[raw]
                elif single_goal_id:
                    oc["parent_id"] = single_goal_id

            # ---- Resolve parents for Outputs
            single_outcome_id = st.session_state.outcomes[0]["id"] if len(st.session_state.outcomes) == 1 else None
            for out in st.session_state.outputs:
                raw = pending_output_parent_ref.get(out["id"], "")
                if raw in outcomes_by_text:
                    out["parent_id"] = outcomes_by_text[raw]
                elif single_outcome_id:
                    out["parent_id"] = single_outcome_id

            # ---- KPI Matrix (preferred in hybrid model)
            if "KPI Matrix" in xls.sheet_names:
                kdf = pd.read_excel(xls, sheet_name="KPI Matrix")


                def _clean(v):
                    try:
                        import pandas as pd
                        if pd.isna(v):
                            return ""
                    except Exception:
                        pass
                    return str(v).strip()


                # Build quick lookup maps by name
                outputs_by_name = {(o.get("name") or "").strip(): o["id"] for o in st.session_state.outputs}

                for _, row in kdf.iterrows():
                    parent_label = _clean(row.get("Parent (label)", ""))  # e.g., "Output 1 — Title"

                    # Always Output in hybrid:
                    plevel = "Output"
                    tail = parent_label.split("—", 1)[1].strip() if "—" in parent_label else parent_label
                    parent_id = outputs_by_name.get(tail)
                    if not parent_id:  # fallback: try full label, then plain
                        parent_id = outputs_by_name.get(parent_label) or outputs_by_name.get(tail)

                    # Pull KPI fields
                    kpi_raw = _clean(row.get("KPI", ""))
                    # strip any leading "KPI x.y — "
                    kpi_text = re.sub(r'^\s*KPI\s+[\w\.\-]+\s*[—:\-]\s*', '', kpi_raw).strip()

                    baseline = _clean(row.get("Baseline", ""))
                    target = _clean(row.get("Target", ""))
                    sd = parse_date_like(row.get("Start Date", ""))
                    ed = parse_date_like(row.get("End Date", ""))
                    linked = _clean(row.get("Linked to Payment", "")).lower() in ("yes", "y", "true", "1")
                    mov = _clean(row.get("Means of Verification", ""))

                    # Only add if we found a parent
                    if parent_id:
                        st.session_state.kpis.append({
                            "id": generate_id(),
                            "level": "KPI",
                            "name": kpi_text,
                            "parent_id": parent_id,
                            "baseline": baseline,
                            "target": target,
                            "start_date": sd,
                            "end_date": ed,
                            "linked_payment": bool(linked),
                            "mov": mov,
                        })

            # ---- Workplan (supports both "rich" and "simple" exports)
            if "Workplan" in xls.sheet_names:
                wdf = pd.read_excel(xls, sheet_name="Workplan")
                # normalize headers
                wdf.columns = [str(c).strip() for c in wdf.columns]


                def _s(v):
                    try:
                        import pandas as pd
                        if pd.isna(v):
                            return ""
                    except Exception:
                        pass
                    return str(v).strip()


                def _split_csv(s):
                    return [t.strip() for t in (_s(s).split(",") if _s(s) else []) if t.strip()]

                # lookups
                outputs_by_name = {(o.get("name") or "").strip(): o["id"] for o in st.session_state.outputs}
                kpis_by_name = {(k.get("name") or "").strip(): k["id"] for k in st.session_state.kpis}
                kpi_id_set = set(kpis_by_name.values())

                # detect "rich" vs "simple" format
                rich = {"Activity ID", "Output", "Activity", "Owner", "Start", "End", "Status", "% complete",
                        "Linked KPIs", "Milestones", "Notes", "Dependencies"}.issubset(set(wdf.columns))

                st.session_state.workplan = []  # reset before loading

                if rich:
                    for _, row in wdf.iterrows():
                        act_id = _s(row.get("Activity ID")) or generate_id()
                        out_name = _s(row.get("Output"))
                        out_id = outputs_by_name.get(out_name)

                        # map KPI names to ids (export wrote names)
                        kpi_names = _split_csv(row.get("Linked KPIs"))
                        kpi_ids = [kpis_by_name.get(n) for n in kpi_names if kpis_by_name.get(n)]

                        # dependencies: prefer IDs if they exist, else leave empty
                        deps = _split_csv(row.get("Dependencies"))
                        dep_ids = [d for d in deps if d in kpi_id_set or len(d) >= 6]  # crude but safe; keep as-is

                        st.session_state.workplan.append({
                            "id": act_id,
                            "output_id": out_id,
                            "name": _s(row.get("Activity")),
                            "owner": _s(row.get("Owner")),
                            "start": parse_date_like(row.get("Start")),
                            "end": parse_date_like(row.get("End")),
                            "status": _s(row.get("Status") or "planned"),
                            "progress": int(float(row.get("% complete") or 0)),
                            "kpi_ids": kpi_ids,
                            "milestones": [m.strip() for m in _s(row.get("Milestones")).split("|") if m.strip()],
                            "dependencies": dep_ids,
                            "notes": _s(row.get("Notes")),
                        })
                else:
                    # SIMPLE legacy format: Activity | Owner | Start Date | End Date | Milestone
                    # No Output/KPI info in this shape; if there is exactly one Output, attach to it.
                    only_output_id = st.session_state.outputs[0]["id"] if len(st.session_state.outputs) == 1 else None
                    for _, row in wdf.iterrows():
                        st.session_state.workplan.append({
                            "id": generate_id(),
                            "output_id": only_output_id,  # None if multiple outputs; user can reassign in UI
                            "name": _s(row.get("Activity")),
                            "owner": _s(row.get("Owner")),
                            "start": parse_date_like(row.get("Start Date")),
                            "end": parse_date_like(row.get("End Date")),
                            "status": "planned",
                            "progress": 0,
                            "kpi_ids": [],
                            "milestones": [_s(row.get("Milestone"))] if _s(row.get("Milestone")) else [],
                            "dependencies": [],
                            "notes": "",
                        })

            # ---- Budget import (optional sheet)
            if "Budget" in xls.sheet_names:
                bdf = pd.read_excel(xls, sheet_name="Budget")
                bdf.columns = [str(c).strip() for c in bdf.columns]

                # Current Outputs (created from Summary) -> build name map
                outputs_by_name = {(o.get("name") or "").strip(): o["id"] for o in st.session_state.outputs}
                current_ids = {o["id"] for o in st.session_state.outputs}

                imported = []
                for _, r in bdf.iterrows():
                    # 1) Prefer Output name if present
                    out_name = str(r.get("Output", "")).strip()
                    out_id = outputs_by_name.get(out_name)

                    # 2) Fallback: try OutputID *only* if it matches a current id
                    if not out_id:
                        raw_id = r.get("OutputID")
                        if raw_id and str(raw_id) in current_ids:
                            out_id = str(raw_id)

                    # Read row values
                    item = str(r.get("Item", "")).strip()
                    cat = str(r.get("Category", "")).strip()
                    unit = str(r.get("Unit", "")).strip()
                    qty = float(r.get("Qty", 0) or 0)
                    uc = float(r.get("Unit Cost", 0) or 0)
                    cur = str(r.get("Currency", "USD")).strip() or "USD"
                    tot = float(r.get("Total", qty * uc) or (qty * uc))

                    # Only keep rows we can link to an Output and that have an item
                    if out_id and item:
                        imported.append([out_id, item, cat, unit, qty, uc, cur, tot])

                if imported:
                    st.session_state.budget = imported

            # --- Enforce single Goal and single Outcome after import ---
            # Keep the first Goal only
            if len(st.session_state.impacts) > 1:
                keep_goal_id = st.session_state.impacts[0]["id"]
                st.session_state.impacts = [st.session_state.impacts[0]]
            else:
                keep_goal_id = st.session_state.impacts[0]["id"] if st.session_state.impacts else None

            # Keep the first Outcome only; reattach all Outputs to it
            if len(st.session_state.outcomes) > 1:
                keep_outcome = st.session_state.outcomes[0]
                keep_outcome_id = keep_outcome["id"]
                st.session_state.outcomes = [keep_outcome]
                for o in st.session_state.outputs:
                    o["parent_id"] = keep_outcome_id   # << this is the line you asked about
            elif st.session_state.outcomes:
                keep_outcome_id = st.session_state.outcomes[0]["id"]
            else:
                keep_outcome_id = None

            # Ensure the (single) Outcome points at the (single) Goal, if both exist
            if keep_goal_id and keep_outcome_id:
                st.session_state.outcomes[0]["parent_id"] = keep_goal_id

            # --- Import Identification sheet (if present) and update ID page state ---
            if "Identification" in xls.sheet_names:
                id_df = pd.read_excel(xls, sheet_name="Identification")
                try:
                    kv = {str(r["Field"]).strip(): str(r["Value"]) if not pd.isna(r["Value"]) else ""
                          for _, r in id_df.iterrows()}
                except Exception:
                    kv = {}

                def _g(field):  # helper to get a string from the kv map
                    return (kv.get(field, "") or "").strip()

                id_info = st.session_state.get("id_info", {}) or {}
                id_info.update({
                    "title": _g("Project title"),
                    "pi_name": _g("Principal Investigator (PI) name"),
                    "pi_email": _g("PI email"),
                    "institution": _g("Institution / Organization"),
                    "start_date": parse_date_like(kv.get("Project start date", "")) or id_info.get("start_date"),
                    "end_date": parse_date_like(kv.get("Project end date", "")) or id_info.get("end_date"),
                    "contact_name": _g("Contact person (optional)"),
                    "contact_email": _g("Contact email"),
                    "contact_phone": _g("Contact phone"),
                })
                st.session_state.id_info = id_info

                # also prime the live widget keys so the inputs show the imported values immediately
                st.session_state["id_title"]        = id_info["title"]
                st.session_state["id_pi_name"]      = id_info["pi_name"]
                st.session_state["id_pi_email"]     = id_info["pi_email"]
                st.session_state["id_institution"]  = id_info["institution"]
                st.session_state["id_start_date"]   = id_info["start_date"]
                st.session_state["id_end_date"]     = id_info["end_date"]
                st.session_state["id_contact_name"] = id_info["contact_name"]
                st.session_state["id_contact_email"]= id_info["contact_email"]
                st.session_state["id_contact_phone"]= id_info["contact_phone"]

            # Remember we loaded this file content; prevents re-import on button clicks
            st.session_state["_resume_file_sig"] = file_sig
            tabs[0].success("✅ Previous submission loaded into session.")
            st.rerun()

        # else: same file uploaded again → skip re-import so edit/delete works
    except Exception as e:
        tabs[0].error(f"Could not parse uploaded Excel: {e}")

# ===== TAB 2: Identification =====
with tabs[1]:
    st.header("🪪 Project Identification")

    # defaults
    if "id_info" not in st.session_state:
        st.session_state.id_info = {
            "title": "", "pi_name": "", "pi_email": "", "institution": "",
            "start_date": None, "end_date": None,
            "contact_name": "", "contact_email": "", "contact_phone": ""
        }

    # ensure widget keys exist (so inputs are persistent & can be set by the importer)
    for k, v in [
        ("id_title", st.session_state.id_info["title"]),
        ("id_pi_name", st.session_state.id_info["pi_name"]),
        ("id_pi_email", st.session_state.id_info["pi_email"]),
        ("id_institution", st.session_state.id_info["institution"]),
        ("id_start_date", st.session_state.id_info["start_date"]),
        ("id_end_date", st.session_state.id_info["end_date"]),
        ("id_contact_name", st.session_state.id_info["contact_name"]),
        ("id_contact_email", st.session_state.id_info["contact_email"]),
        ("id_contact_phone", st.session_state.id_info["contact_phone"]),
    ]:
        if k not in st.session_state:
            st.session_state[k] = v

    c1, c2 = st.columns(2)
    with c1:
        st.session_state.id_info["title"] = st.text_input("Project title*", key="id_title")
        st.session_state.id_info["pi_name"] = st.text_input("Principal Investigator (PI) name*", key="id_pi_name")
        st.session_state.id_info["pi_email"] = st.text_input("PI email*", key="id_pi_email")
        st.session_state.id_info["institution"] = st.text_input("Institution / Organization*", key="id_institution")
    with c2:
        sd = st.date_input("Project start date*", key="id_start_date")
        if sd:
            st.caption(f"Selected: {fmt_dd_mmm_yyyy(sd)}")  # show DD/Mon/YYYY preview

        ed = st.date_input("Project end date*", key="id_end_date")
        if ed:
            st.caption(f"Selected: {fmt_dd_mmm_yyyy(ed)}")  # show DD/Mon/YYYY preview

        # keep canonical values in id_info for export/validation
        st.session_state.id_info["start_date"] = sd
        st.session_state.id_info["end_date"] = ed
        with st.expander("More contact details (optional)"):
            st.session_state.id_info["contact_name"] = st.text_input("Contact person (if different from PI)", key="id_contact_name")
            st.session_state.id_info["contact_email"] = st.text_input("Contact email", key="id_contact_email")
            st.session_state.id_info["contact_phone"] = st.text_input("Contact phone", key="id_contact_phone")

    # inline validation (no button)
    errs = []
    ii = st.session_state.id_info
    if not ii["title"].strip():       errs.append("Project title is required.")
    if not ii["pi_name"].strip():     errs.append("PI name is required.")
    if not ii["institution"].strip(): errs.append("Institution is required.")
    if not ii["pi_email"].strip() or "@" not in ii["pi_email"]: errs.append("Valid PI email is required.")
    if ii["start_date"] and ii["end_date"] and ii["start_date"] > ii["end_date"]:
        errs.append("Project start date must be on or before the end date.")

    # show errors only after the user has typed something in any required field
    touched = any([
        st.session_state["id_title"].strip(),
        st.session_state["id_pi_name"].strip(),
        st.session_state["id_pi_email"].strip(),
        st.session_state["id_institution"].strip(),
    ])
    if touched:
        for e in errs:
            st.error(e)

    # --- Read-only summary (live)
    # Budget total (computed from detailed budget)
    def _sum_budget():
        total = 0.0
        for row in st.session_state.get("budget", []):
            try:
                # your budget rows are [OutputID, Item, Category, Unit, Qty, Unit Cost, Currency, Total]
                total += float(row[7])
            except Exception:
                pass
        return total

    budget_total = _sum_budget()

    # Live counts
    outputs_count = len(st.session_state.get("outputs", []))
    # activities_count = len(st.session_state.get("activities", []))
    kpis_count = len(st.session_state.get("kpis", []))

    st.markdown("### Summary")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Funding requested (from Budget)**")
        st.markdown(f"<div style='font-size:1.5em; font-weight:600;'>{fmt_money(budget_total)}</div>",
                    unsafe_allow_html=True)
    with c2:
        st.markdown("**Logframe indicators**")
        st.markdown(
            f"""
            <div style="display:flex; gap:10px; align-items:center;">
              <div style="background:#eef2ff;border:1px solid #dbe2ff;border-radius:999px;padding:6px 10px;">Outputs: <b>{outputs_count}</b></div>
              <div style="background:#f7f7f9;border:1px solid #e6e6e6;border-radius:999px;padding:6px 10px;">KPIs: <b>{kpis_count}</b></div>
            </div>
            """,
            unsafe_allow_html=True
        )

    # Cross-check note area (non-blocking warnings placeholder)
    warnings = []
    # (Optional) You can later add KPI due-date vs project dates checks here and append messages to `warnings`.
    if warnings:
        for w in warnings:
            st.warning(w)


# ===== TAB 3: Logframe =====
tabs[2].header("📊 Build Your Logframe")
inject_logframe_css()
# --- numbering for labels shown in UI and preview ---
out_nums, kpi_nums = compute_numbers()   # outputs -> 'n', KPIs -> 'n.p'

# --- Add forms ---
with tabs[2].expander("➕ Add Goal"):
    if len(st.session_state.impacts) >= 1:
        st.info("Only one Goal is allowed. Edit the existing Goal in the preview below.")
    else:
        with st.form("goal_form"):
            goal_text = st.text_area("Goal (single, high-level statement)")
            if st.form_submit_button("Add Goal") and goal_text.strip():
                st.session_state.impacts.append({"id": generate_id(), "level": "Goal", "name": goal_text.strip()})

with tabs[2].expander("➕ Add Outcome"):
    if not st.session_state.impacts:
        st.warning("Add the Goal first.")
    elif len(st.session_state.outcomes) >= 1:
        st.info("Only one Outcome is allowed. Edit the existing Outcome in the preview below.")
    else:
        with st.form("outcome_form"):
            outcome_text = st.text_area("Outcome (statement)")
            # since there is only one goal, no need to pick it; link to the single goal
            linked_goal_id = st.session_state.impacts[0]["id"]
            if st.form_submit_button("Add Outcome") and outcome_text.strip():
                st.session_state.outcomes.append(
                    {"id": generate_id(), "level": "Outcome", "name": outcome_text.strip(), "parent_id": linked_goal_id}
                )

with tabs[2].expander("➕ Add Output"):
    if not st.session_state.outcomes:
        tabs[2].warning("Add the Outcome first.")
    else:
        with st.form("output_form"):
            output_title = st.text_input("Output title (e.g., 'Output 1')")
            output_assumptions = st.text_area("Key Assumptions (optional)")
            if st.form_submit_button("Add Output") and output_title.strip():
                linked_outcome_id = st.session_state.outcomes[0]["id"]
                st.session_state.outputs.append(
                    {
                        "id": generate_id(),
                        "level": "Output",
                        "name": output_title.strip(),
                        "parent_id": linked_outcome_id,
                        "assumptions": output_assumptions.strip(),
                    }
                )

with tabs[2].expander("➕ Add KPI"):
    if not st.session_state.outputs:
        tabs[2].warning("Add an Output first.")
    else:
        with st.form("kpi_form"):
            parent = st.selectbox(
                "Parent Output",
                st.session_state.outputs,
                format_func=lambda o: f"Output {out_nums.get(o['id'],'?')} — {o.get('name','Output')}"
            )
            kpi_text = st.text_area("KPI*")
            baseline = st.text_input("Baseline")
            target   = st.text_input("Target")
            c1, c2 = st.columns(2)
            with c1:
                start_date = st.date_input("Start date")
            with c2:
                end_date   = st.date_input("End date")
            payment_linked = st.checkbox("Linked to Payment (optional)")
            mov = st.text_area("Means of Verification")

            if st.form_submit_button("Add KPI") and kpi_text.strip():
                st.session_state.kpis.append({
                    "id": generate_id(),
                    "level": "KPI",
                    "name": strip_label_prefix(kpi_text.strip(), "KPI"),
                    "parent_id": parent["id"],
                    "parent_level": "Output",   # <-- fixed
                    "baseline": baseline.strip(),
                    "target": target.strip(),
                    "start_date": start_date,
                    "end_date": end_date,
                    "linked_payment": bool(payment_linked),
                    "mov": mov.strip(),
                })
                st.rerun()

# ---- View helpers (card layout, compact-aware) ----
def view_goal(g):
    return f"##### 🟦 **Goal:** {g.get('name','')}"

def view_outcome(o):
    return f"##### 🟪 **Outcome:** {o.get('name','')}"

def view_output(out):
    num   = out_nums.get(out["id"], "?")
    title = out.get("name", "Output")

    # header line (keeps your green card style via view_logframe_element)
    header_html = (
        f"<div class='lf-out-header'><strong>Output {num}:</strong> {title}</div>"
    )

    # assumptions -> bullet list (one bullet per line the user typed)
    ass = (out.get("assumptions") or "").strip()
    ass_html = ""
    if ass:
        # strip any leading "-" or "•", ignore empty lines
        items = [
            re.sub(r"^[\-\u2022]\s*", "", ln).strip()
            for ln in ass.splitlines()
            if ln.strip()
        ]
        if items:
            lis = "".join(f"<li>{html.escape(x)}</li>" for x in items)
            ass_html = (
                "<div class='lf-ass'>"
                "<div class='lf-ass-heading'> <b> Key Assumptions </b> </div>"
                f"<ul class='lf-ass-list'>{lis}</ul>"
                "</div>"
            )

    # wrap in the green card
    return view_logframe_element(header_html + ass_html, kind="output")

def view_output_header(out):
    num   = out_nums.get(out["id"], "?")
    title = escape(out.get("name", "Output"))
    header_html = f"<div class='lf-out-header'><strong>Output {num}:</strong> {title}</div>"
    return view_logframe_element(header_html, kind="output")

def view_kpi(k):
    num  = kpi_nums.get(k["id"], "?")
    name = k.get("name", "")

    bp  = (k.get("baseline") or "").strip()
    tg  = (k.get("target") or "").strip()
    sd  = fmt_dd_mmm_yyyy(k.get("start_date")) or "—"
    ed  = fmt_dd_mmm_yyyy(k.get("end_date")) or "—"
    mov = (k.get("mov") or "").strip()

    chip = (
        "<span class='chip green'>Payment-linked</span>"
        if k.get("linked_payment")
        else "<span class='chip'>Not payment-linked</span>"
    )

    # Title
    header = f"<div class='lf-kpi-title'>KPI {num}: {name}</div>"

    # Lines in the exact order you want
    lines = []
    if bp:
        lines.append(f"<div class='lf-line'><b>Baseline:</b> {bp}</div>")
    if tg:
        lines.append(f"<div class='lf-line'><b>Target:</b> {tg}</div>")
    # dates on one line (use em dash if you prefer: &mdash;)
    if (sd != "—") or (ed != "—"):
        lines.append(
            f"<div class='lf-line'><b>Start date:</b> {sd} &nbsp;&nbsp;•&nbsp;&nbsp; "
            f"<b>End date:</b> {ed}</div>"
        )
    if mov:
        lines.append(f"<div class='lf-line'><b>Means of Verification:</b> {mov}</div>")

    # chip gets its own line, independent from dates
    lines.append(f"<div class='lf-line'>{chip}</div>")

    inner = header + "".join(lines)
    return view_logframe_element(inner, kind="kpi")

def view_activity(a: dict, act_label: str, id_to_output: dict, id_to_kpi: dict) -> str:
    """
    Render one activity as a card.
    - a: activity dict from st.session_state.workplan
    - act_label: precomputed label like "1.2"
    - id_to_output: {output_id -> output name}
    - id_to_kpi:    {kpi_id -> kpi name}
    """
    out_name = id_to_output.get(a.get("output_id"), "(unassigned)")
    title    = f"Activity {escape(act_label)} — {escape(a.get('name',''))}"
    owner    = escape(a.get("owner","") or "—")
    sd       = fmt_dd_mmm_yyyy(a.get("start")) or "—"
    ed       = fmt_dd_mmm_yyyy(a.get("end"))   or "—"
    status   = escape(a.get("status","planned"))
    prog     = f"{int(a.get('progress',0))}%"
    kpis_txt = ", ".join(escape(id_to_kpi.get(kid,"")) for kid in (a.get("kpi_ids") or [])) or "—"
    miles    = " | ".join(escape(m) for m in (a.get("milestones") or [])) or "—"
    notes    = escape(a.get("notes","") or "")

    title_html = f"<div class='lf-activity-title'>{title}</div>"
    rows = [
        ("Output", out_name),
        ("Owner", owner),
        ("Start date", sd),
        ("End date", ed),
        ("Status", status),
        ("% complete", prog),
        ("Linked KPIs", kpis_txt),
        ("Milestones", miles),
        ("Notes", notes),  # the card helper will skip empty values if you prefer
    ]

    # Build rows (simple label/value pairs). You can skip empty ones here if desired.
    body = "".join(
        f"<div class='lf-line'><b>{escape(lbl)}:</b> {val}</div>"
        for (lbl, val) in rows
        if (val is not None and str(val).strip() != "")
    )

    # Use your generic card wrapper (orange indicators style)
    return view_logframe_element(title_html + body, kind="activity")

# --- Inline preview with Edit / Delete buttons (refactored, card layout) ---
with tabs[2]:
    st.markdown("---")
    st.subheader("Current Logframe (preview) — click ✏️ to edit, 🗑️ to delete")

    for g in st.session_state.get("impacts", []):
        render_editable_item(
            item=g,
            list_name="impacts",
            edit_flag_key="edit_goal",
            view_md_func=view_goal,
            default_label="Goal",
            on_delete=lambda _id=g["id"]: (delete_cascade(goal_id=_id), st.rerun()),
            key_prefix="lf"
        )

        outcomes_here = [o for o in st.session_state.get("outcomes", []) if o.get("parent_id") == g["id"]]
        for oc in outcomes_here:
            render_editable_item(
                item=oc,
                list_name="outcomes",
                edit_flag_key="edit_outcome",
                view_md_func=view_outcome,
                default_label="Outcome",
                on_delete=lambda _id=oc["id"]: (delete_cascade(outcome_id=_id), st.rerun()),
                key_prefix="lf"
            )

            outs_here = [o for o in st.session_state.get("outputs", []) if o.get("parent_id") == oc["id"]]
            for out in outs_here:
                with st.container():  # now this container lives inside the Logframe tab
                    render_editable_item(
                        item=out,
                        list_name="outputs",
                        edit_flag_key="edit_output",
                        view_md_func=view_output,
                        fields=[
                            ("name", st.text_input, "Output title"),
                            ("assumptions", st.text_area, "Key Assumptions"),
                        ],
                        on_delete=lambda _id=out["id"]: (delete_cascade(output_id=_id), st.rerun()),
                        key_prefix="lf"
                    )

                    k_children = [k for k in st.session_state.get("kpis", [])
                                  if (k.get("parent_id") == out["id"])]  # parent_level check optional if you dropped outcome-level KPIs
                    for k in k_children:
                        render_editable_item(
                            item=k, list_name="kpis", edit_flag_key="edit_kpi",
                            view_md_func=view_kpi,
                            fields=[
                                ("name", st.text_area, "KPI"),
                                ("baseline", st.text_input, "Baseline"),
                                ("target", st.text_input, "Target"),
                                ("start_date",
                                 lambda label, value, key: st.date_input(label, value=value or date.today(), key=key),
                                 "Start date"),
                                ("end_date",
                                 lambda label, value, key: st.date_input(label, value=value or date.today(), key=key),
                                 "End date"),
                                ("linked_payment",
                                 lambda label, value, key: st.checkbox(label, value=bool(value), key=key),
                                 "Linked to Payment"),
                                ("mov", st.text_area, "Means of Verification"),
                            ],
                            on_delete=lambda _id=k["id"]: (
                                setattr(st.session_state, "kpis", [x for x in st.session_state.kpis if x["id"] != _id]),
                                st.rerun()
                            ),
                            key_prefix="lf"
                        )

# ===== TAB 4: Workplan =====
with tabs[3]:
    st.header("📆 Workplan")
    with st.expander("➕ Add Activity"):
        with st.form("workplan_form_v2"):
            # Required: link to Output
            output_parent = st.selectbox(
                "Linked Output*",
                st.session_state.outputs,
                format_func=lambda x: x.get("name") or "Output"
            )

            # Optional: link to KPI(s) under that Output
            output_id = output_parent["id"] if output_parent else None
            kpis_for_output = [k for k in st.session_state.kpis if k.get("parent_id") == output_id]
            kpi_links = st.multiselect(
                "Linked KPI(s) (optional)",
                kpis_for_output,
                format_func=lambda k: f"{k.get('name','')}"
            )

            name = st.text_input("Activity*")
            owner = st.text_input("Responsible person/institution*")
            c1, c2 = st.columns(2)
            with c1:
                start = st.date_input("Start date*")
            with c2:
                end = st.date_input("End date*")

            status = st.selectbox("Status", ["planned", "in_progress", "completed", "cancelled"], index=0)
            progress = st.slider("% complete", 0, 100, 0)
            milestones = st.text_area("Milestones / deliverables (optional, one per line)")
            notes = st.text_area("Notes (optional)")

            # Dependencies: choose among existing activities (by name)
            existing_acts = st.session_state.get("workplan", [])
            deps = st.multiselect(
                "Depends on (optional)",
                existing_acts,
                format_func=lambda a: a.get("name") if isinstance(a, dict) else (a[0] if a else "")
            )

            submitted = st.form_submit_button("Add to Workplan")
            if submitted and name.strip() and owner.strip() and output_parent and start and end and start <= end:
                st.session_state.workplan.append({
                    "id": generate_id(),
                    "output_id": output_id,
                    "name": name.strip(),
                    "kpi_ids": [k["id"] for k in kpi_links],
                    "owner": owner.strip(),
                    "start": start,
                    "end": end,
                    "status": status,
                    "progress": int(progress),
                    "milestones": [m.strip() for m in milestones.splitlines() if m.strip()],
                    "dependencies": [ (d.get("id") if isinstance(d, dict) else None) for d in deps ],
                    "notes": notes.strip() if notes else ""
                })
                st.rerun()
            elif submitted:
                st.warning("Please fill required fields (Output, Activity, Owner, Start≤End).")

    # --- Card view (optional: put this below or instead of the table) ---
    out_nums, kpi_nums, act_nums = compute_numbers(include_activities=True)
    id_to_output = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}
    id_to_kpi    = {k["id"]: (k.get("name") or "")       for k in st.session_state.kpis}

    for oc in st.session_state.outcomes:
        outs_here = [o for o in st.session_state.outputs if o.get("parent_id") == oc["id"]]
        for out in outs_here:
            # green Output header card
            st.markdown(view_output_header(out), unsafe_allow_html=True)

            # orange Activity cards (with edit/delete)
            acts_here = [a for a in st.session_state.workplan if a.get("output_id") == out["id"]]
            for a in acts_here:
                label = act_nums.get(a["id"], "?")
                # Edit mode?
                if st.session_state.get("edit_activity") == a["id"]:
                    e1, e2, e3 = st.columns([0.90, 0.05, 0.05])
                    with e1:
                        new_name = st.text_input("Activity", value=a.get("name", ""), key=f"a_name_{a['id']}")
                        new_owner = st.text_input("Owner", value=a.get("owner", ""), key=f"a_owner_{a['id']}")
                        cA, cB = st.columns(2)
                        with cA:
                            new_start = st.date_input("Start date", value=a.get("start"), key=f"a_start_{a['id']}")
                        with cB:
                            new_end = st.date_input("End date", value=a.get("end"), key=f"a_end_{a['id']}")
                        new_status = st.selectbox(
                            "Status", ["planned", "in_progress", "completed", "cancelled"],
                            index=["planned", "in_progress", "completed", "cancelled"].index(
                                a.get("status", "planned")),
                            key=f"a_status_{a['id']}"
                        )
                        new_prog = st.slider("% complete", 0, 100, int(a.get("progress", 0)), key=f"a_prog_{a['id']}")
                        new_notes = st.text_area("Notes", value=a.get("notes", ""), key=f"a_notes_{a['id']}")
                    if e2.button("💾", key=f"a_save_{a['id']}"):
                        idx = _find_by_id(st.session_state.workplan, a["id"])
                        if idx is not None:
                            st.session_state.workplan[idx].update({
                                "name": new_name.strip(),
                                "owner": new_owner.strip(),
                                "start": new_start,
                                "end": new_end,
                                "status": new_status,
                                "progress": int(new_prog),
                                "notes": new_notes.strip(),
                            })
                        st.session_state["edit_activity"] = None
                        st.rerun()
                    if e3.button("✖️", key=f"a_cancel_{a['id']}"):
                        st.session_state["edit_activity"] = None
                        st.rerun()
                else:
                    v1, v2, v3 = st.columns([0.90, 0.05, 0.05])
                    v1.markdown(
                        view_activity_readonly(a, label, id_to_output, id_to_kpi),
                        unsafe_allow_html=True
                    )
                    if v2.button("✏️", key=f"a_edit_{a['id']}"):
                        st.session_state["edit_activity"] = a["id"]
                        st.rerun()
                    if v3.button("🗑️", key=f"a_del_{a['id']}"):
                        st.session_state.workplan = [x for x in st.session_state.workplan if x["id"] != a["id"]]
                        st.rerun()

# ===== TAB 5: Budget =====
with tabs[4]:
    st.header("💵 Define Budget\nEnter amounts in USD")

    # ---------- Add new budget item ----------
    with st.expander("➕ Add Budget Item"):
        with st.form("budget_form"):
            if not st.session_state.outputs:
                st.warning("Add an Output first (in the Logframe tab) before adding budget lines.")
                st.form_submit_button("Add to Budget", disabled=True)
            else:
                output_parent = st.selectbox(
                    "Linked Output*",
                    st.session_state.outputs,
                    format_func=lambda x: x.get("name") or "Output",
                    index=0
                )
                item      = st.text_input("Item*")
                category  = st.selectbox("Category", ["Personnel","Supplies","Travel","Equipment","Services","Other"])
                unit      = st.text_input("Unit (e.g., day, set, PM)")
                quantity  = st.number_input("Quantity", min_value=0.0, value=1.0)
                unit_cost = st.number_input("Unit Cost", min_value=0.0, value=0.0)
                currency  = st.text_input("Currency (ISO 4217)", value="USD")

                total = round(float(quantity) * float(unit_cost), 2)

                if st.form_submit_button("Add to Budget"):
                    if not item.strip():
                        st.warning("Item is required.")
                    else:
                        # canonical 8-column schema
                        st.session_state.budget.append([
                            output_parent["id"],         # OutputID
                            item.strip(),                # Item
                            category,                    # Category
                            unit.strip(),                # Unit
                            float(quantity),             # Qty
                            float(unit_cost),            # Unit Cost
                            currency.strip(),            # Currency
                            float(total)                 # Total
                        ])
                        st.rerun()

    # ---------- Grouped by Output, compact blue rows ----------
    id_to_output_name = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}

    def render_budget_row_inline(row):
        """Return compact inline HTML for one budget row."""
        out_id, item, cat, unit, qty, uc, cur, tot = row
        return (
            "<div class='lf-budget-row'>"
            "  <div class='lf-budget-cells'>"
            f"    <div class='lf-budget-cell'><span class='lf-budget-chip'>{escape(item or '—')}</span></div>"
            f"    <div class='lf-budget-cell'>{escape(cat or '—')}</div>"
            f"    <div class='lf-budget-cell'>{escape(unit or '—')}</div>"
            f"    <div class='lf-budget-cell'>Qty: {qty or 0:g}</div>"
            f"    <div class='lf-budget-cell lf-budget-money'>Unit: {fmt_money(uc or 0)}</div>"
            f"    <div class='lf-budget-cell'>{escape(cur or '—')}</div>"
            f"    <div class='lf-budget-cell lf-budget-money'><b>{fmt_money(tot or 0)}</b></div>"
            "  </div>"
            "</div>"
        )

    grand_total = 0.0

    # outputs in the same order as in Logframe
    out_nums, _ = compute_numbers()
    outputs_sorted = sorted(st.session_state.outputs, key=lambda o: int(out_nums.get(o["id"], "9999").split('.')[0]))

    for out in outputs_sorted:
        # Header like Workplan
        st.markdown(view_output_header(out), unsafe_allow_html=True)

        # Items for this output (keep order they were added)
        rows_here = [(idx, r) for idx, r in enumerate(st.session_state.budget) if r[0] == out["id"]]
        if not rows_here:
            st.info("No budget items for this output.")
            continue

        subtotal = 0.0
        for idx, r in rows_here:
            out_id, item, cat, unit, qty, uc, cur, tot = r
            subtotal += float(tot or 0.0)
            grand_total += float(tot or 0.0)

            # Unique suffix for keys (prevents duplicate keys even if order changes)
            row_uid = f"{idx}_{out_id}_{hash(item)}_{int(qty)}_{int(uc)}"

            if st.session_state.get("edit_budget_row") == idx:
                # ----- inline edit mode (row replaced by small form) -----
                e1, e2, e3 = st.columns([0.90, 0.05, 0.05])
                with e1:
                    out_sel = st.selectbox(
                        "Output", st.session_state.outputs,
                        format_func=lambda x: x.get("name") or "Output",
                        index=next((j for j,o in enumerate(st.session_state.outputs) if o["id"] == out_id), 0),
                        key=f"b_out_{row_uid}"
                    )
                    new_item = st.text_input("Item", value=item, key=f"b_item_{row_uid}")
                    new_cat  = st.selectbox(
                        "Category", ["Personnel","Supplies","Travel","Equipment","Services","Other"],
                        index=(["Personnel","Supplies","Travel","Equipment","Services","Other"].index(cat)
                               if cat in ["Personnel","Supplies","Travel","Equipment","Services","Other"] else 0),
                        key=f"b_cat_{row_uid}"
                    )
                    cols = st.columns(3)
                    with cols[0]:
                        new_unit = st.text_input("Unit", value=unit, key=f"b_unit_{row_uid}")
                    with cols[1]:
                        new_qty  = st.number_input("Qty", min_value=0.0, value=float(qty or 0), key=f"b_qty_{row_uid}")
                    with cols[2]:
                        new_uc   = st.number_input("Unit Cost", min_value=0.0, value=float(uc or 0), key=f"b_uc_{row_uid}")
                    new_cur = st.text_input("Currency", value=cur or "USD", key=f"b_cur_{row_uid}")
                    new_tot = round(float(new_qty) * float(new_uc), 2)
                    st.caption(f"New total: {fmt_money(new_tot)}")

                if e2.button("💾", key=f"b_save_{row_uid}"):
                    st.session_state.budget[idx] = [
                        out_sel["id"], new_item.strip(), new_cat, new_unit.strip(),
                        float(new_qty), float(new_uc), new_cur.strip(), float(new_tot)
                    ]
                    st.session_state["edit_budget_row"] = None
                    st.rerun()

                if e3.button("✖️", key=f"b_cancel_{row_uid}"):
                    st.session_state["edit_budget_row"] = None
                    st.rerun()

            else:
                # ----- view mode: compact row + buttons -----
                c1, c2, c3 = st.columns([0.92, 0.04, 0.04])
                c1.markdown(render_budget_row_inline(r), unsafe_allow_html=True)
                if c2.button("✏️", key=f"b_edit_{row_uid}"):
                    st.session_state["edit_budget_row"] = idx
                    st.rerun()
                if c3.button("🗑️", key=f"b_del_{row_uid}"):
                    del st.session_state.budget[idx]
                    st.rerun()

        # Subtotal under the group
        st.markdown(f"<div class='lf-subtotal'>Subtotal: {fmt_money(subtotal)}</div>", unsafe_allow_html=True)

    # Grand total at the bottom
    st.markdown(f"<div class='lf-grandtotal'>Total: {fmt_money(grand_total)}</div>", unsafe_allow_html=True)


# ===== TAB 6: Export =====
tabs[5].header("📤 Export Your Application")
if tabs[5].button("Generate Excel File"):
    wb = Workbook()
    if "Sheet" in wb.sheetnames:
        wb.remove(wb["Sheet"])

    # --- Sheet 0: Identification (Project ID page) ---
    def _sum_budget_for_export():
        total = 0.0
        for r in st.session_state.get("budget", []):
            try:
                total += float(r[7])  # Total at index 7 in current budget schema
            except Exception:
                pass
        return total


    id_info = st.session_state.get("id_info", {}) or {}

    proj_title = id_info.get("title", "")
    pi_name = id_info.get("pi_name", "")
    pi_email = id_info.get("pi_email", "")
    institution = id_info.get("institution", "")
    start_date = id_info.get("start_date", "")
    end_date = id_info.get("end_date", "")
    contact_name = id_info.get("contact_name", "")
    contact_mail = id_info.get("contact_email", "")
    contact_phone = id_info.get("contact_phone", "")

    # live computed
    budget_total = _sum_budget_for_export()
    outputs_count = len(st.session_state.get("outputs", []))
    kpis_count = len(st.session_state.get("kpis", []))

    ws_id = wb.create_sheet("Identification", 0)  # put it first
    ws_id.append(["Field", "Value"])
    ws_id.append(["Project title", proj_title])
    ws_id.append(["Principal Investigator (PI) name", pi_name])
    ws_id.append(["PI email", pi_email])
    ws_id.append(["Institution / Organization", institution])
    ws_id.append(["Project start date", fmt_dd_mmm_yyyy(start_date)])
    ws_id.append(["Project end date", fmt_dd_mmm_yyyy(end_date)])
    ws_id.append(["Contact person (optional)", contact_name])
    ws_id.append(["Contact email", contact_mail])
    ws_id.append(["Contact phone", contact_phone])

    # read-only summary values
    ws_id.append(["Funding requested (from Budget)", f"{budget_total:,.2f}"])
    ws_id.append(["Outputs (count)", outputs_count])
    ws_id.append(["KPIs (count)", kpis_count])

    # Sheet 1: Summary (Goal/Outcome/Output)
    s1 = wb.create_sheet("Summary", 1)  # index 1 = after "Identification"
    s1.append(["Level", "Text / Title", "Parent ID", "Assumptions"])

    # Goal row(s) – no assumptions
    for row in st.session_state.get("impacts", []):
        s1.append([row.get("level", "Goal"), row.get("name", ""), "", ""])

    # Outcome row(s) – no assumptions
    for row in st.session_state.get("outcomes", []):
        s1.append([row.get("level", "Outcome"), row.get("name", ""), row.get("parent_id", ""), ""])

    # Output row(s) – include assumptions
    for row in st.session_state.get("outputs", []):
        s1.append([
            row.get("level", "Output"),
            row.get("name", ""),
            row.get("parent_id", ""),
            row.get("assumptions", "")
        ])

    # Sheet 2: KPI Matrix (Output > KPI)
    s2 = wb.create_sheet("KPI Matrix")
    s2.append([
        "Parent Level",
        "Parent (label)",
        "KPI",
        "Baseline",
        "Target",
        "Start Date",
        "End Date",
        "Linked to Payment",
        "Means of Verification",
    ])

    out_nums, kpi_nums = compute_numbers()
    output_title = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}

    for k in st.session_state.kpis:  # keep order as-is
        pid = k.get("parent_id", "")
        parent_label = f"Output {out_nums.get(pid, '')} — {output_title.get(pid, '')}"
        s2.append([
            "Output",
            parent_label,
            f"KPI {kpi_nums.get(k['id'], '')} — {k.get('name', '')}",
            k.get("baseline", ""),
            k.get("target", ""),
            fmt_dd_mmm_yyyy(k.get("start_date")),
            fmt_dd_mmm_yyyy(k.get("end_date")),
            "Yes" if k.get("linked_payment") else "No",
            k.get("mov", ""),
        ])

    # Workplan (export)
    out_nums, kpi_nums, act_nums = compute_numbers(include_activities=True)
    ws2 = wb.create_sheet("Workplan")
    ws2.append(["Activity ID", "Activity #", "Output", "Activity", "Owner", "Start", "End", "Status", "% complete",
                "Linked KPIs", "Milestones", "Notes", "Dependencies"])
    id_to_output = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}
    id_to_kpi = {k["id"]: (k.get("name") or "") for k in st.session_state.kpis}

    for a in st.session_state.workplan:
        ws2.append([
            a["id"],
            act_nums.get(a["id"], ""),  # ← uses act_nums
            id_to_output.get(a["output_id"], ""),
            a["name"], a["owner"],
            fmt_dd_mmm_yyyy(a["start"]), fmt_dd_mmm_yyyy(a["end"]),
            a["status"], a["progress"],
            ", ".join(id_to_kpi.get(i, "") for i in (a.get("kpi_ids") or [])),
            " | ".join(a.get("milestones") or []),
            a.get("notes", ""),
            ", ".join(filter(None, a.get("dependencies") or [])),
        ])

    # Budget
    # --- Budget (export) ---
    ws3 = wb.create_sheet("Budget")
    ws3.append(["OutputID", "Output", "Item", "Category", "Unit", "Qty", "Unit Cost", "Currency", "Total"])

    # map id -> plain output name (not label)
    id_to_output_name = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}

    for r in st.session_state.budget:
        out_id, item, cat, unit, qty, unit_cost, curr, total = r
        out_name = id_to_output_name.get(out_id, "")
        ws3.append([out_id, out_name, item, cat, unit, qty, unit_cost, curr, total])

    # OPTIONAL Excel number formats (update column indices because we added "Output")
    for row in ws3.iter_rows(min_row=2):
        # columns (1-indexed): 1 OutputID, 2 Output, 3 Item, 4 Category, 5 Unit, 6 Qty, 7 Unit Cost, 8 Currency, 9 Total
        row[5].number_format = '#,##0.00'  # Qty
        row[6].number_format = '#,##0.00'  # Unit Cost
        row[8].number_format = '#,##0.00'  # Total

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    tabs[5].download_button(
        "📥 Download Excel File",
        data=buf,
        file_name="Application_Submission.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

# --- Word export (Logframe as table)
if tabs[5].button("Generate Word Logframe"):
    try:
        word_buf = build_logframe_docx()
        proj_title = (st.session_state.get("id_info", {}) or {}).get("title", "") or "Project"
        safe = re.sub(r"[^A-Za-z0-9]+", "_", proj_title).strip("_") or "Project"
        tabs[5].download_button(
            "📥 Download Word Logframe",
            data=word_buf,
            file_name=f"Logframe_{safe}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ModuleNotFoundError:
        tabs[5].error("`python-docx` is required. Install it with: pip install python-docx")
    except Exception as e:
        tabs[5].error(f"Could not generate the Word logframe: {e}")
